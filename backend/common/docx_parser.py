"""Screenplay .docx preprocessor for N01 pipeline input.

Parses a Word document containing a full screenplay and extracts
structured sections for downstream pipeline processing.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)


@dataclass
class RawScriptInput:
    script_text: str                       # Full screenplay text
    narrative_arc: dict | None = None      # Three-act structure (if present)
    character_presets: list[dict] = field(default_factory=list)
    production_requirements: dict = field(default_factory=dict)
    existing_storyboard: dict | None = None
    project_meta: dict | None = None


def parse_docx(file_path: str) -> RawScriptInput:
    """Parse a .docx screenplay file into RawScriptInput.

    Handles:
    - Full text extraction with paragraph separation
    - Character name detection from dialogue patterns
    - Section header detection for narrative structure
    - Image extraction count for reference
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Screenplay file not found: {file_path}")

    doc = Document(file_path)

    # Extract full text
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    full_text = "\n".join(paragraphs)

    # Detect characters from dialogue patterns (e.g. "张三：" or "JOHN:")
    characters = _extract_characters(paragraphs)

    # Detect narrative structure from headings
    narrative_arc = _extract_narrative_arc(doc)

    # Count embedded images
    image_count = _count_images(doc)

    logger.info(
        "Parsed docx: %d paragraphs, %d characters detected, %d images, %d chars total",
        len(paragraphs),
        len(characters),
        image_count,
        len(full_text),
    )

    return RawScriptInput(
        script_text=full_text,
        narrative_arc=narrative_arc,
        character_presets=characters,
        production_requirements={
            "source_format": "docx",
            "source_file": path.name,
            "paragraph_count": len(paragraphs),
            "image_count": image_count,
            "total_chars": len(full_text),
        },
        existing_storyboard=None,
        project_meta={
            "file_name": path.name,
            "file_size_bytes": path.stat().st_size,
        },
    )


def _extract_characters(paragraphs: list[str]) -> list[dict]:
    """Extract character names from dialogue patterns."""
    # Match patterns like "角色名：" or "角色名:" at line start
    char_pattern = re.compile(r"^([A-Za-z\u4e00-\u9fff]{1,10})[：:]")
    names: dict[str, int] = {}

    for para in paragraphs:
        m = char_pattern.match(para)
        if m:
            name = m.group(1)
            # Filter out common non-character prefixes
            if name not in {"注", "备注", "说明", "提示", "注意", "场景", "画面", "旁白", "字幕"}:
                names[name] = names.get(name, 0) + 1

    # Only keep characters that appear multiple times (likely real characters)
    return [
        {"name": name, "dialogue_count": count}
        for name, count in sorted(names.items(), key=lambda x: -x[1])
        if count >= 2
    ]


def _extract_narrative_arc(doc: Document) -> dict | None:
    """Try to detect three-act structure from heading styles."""
    acts: list[dict] = []

    for para in doc.paragraphs:
        if para.style and para.style.name and "Heading" in para.style.name:
            acts.append({
                "heading": para.text.strip(),
                "level": para.style.name,
            })

    if not acts:
        return None

    return {"headings": acts, "act_count": len(acts)}


def _count_images(doc: Document) -> int:
    """Count embedded images in the document."""
    count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            count += 1
    return count
